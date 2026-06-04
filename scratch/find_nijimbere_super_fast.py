import os
import openpyxl
from docx import Document

workspace_dir = r"c:\Users\Degnon\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy"
target_dirs = [
    os.path.join(workspace_dir, "PROFORMA CHIRURGIE"),
    os.path.join(workspace_dir, "1. Document PC DR GIPSY"),
    os.path.join(workspace_dir, "POSTE DR GIPSY"),
    os.path.join(workspace_dir, "RAPPORT CONS"),
]

def search_excel(path):
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell and "NIJIMBERE" in str(cell).upper():
                        return sheet, cell
    except Exception as e:
        pass
    return None

def search_docx(path):
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            if "NIJIMBERE" in p.text.upper():
                return p.text
        for t in doc.tables:
            for r in t.rows:
                for cell in r.cells:
                    if "NIJIMBERE" in cell.text.upper():
                        return cell.text
    except Exception as e:
        pass
    return None

print("Super fast scan for NIJIMBERE in document folders...")
found = False
for t_dir in target_dirs:
    if not os.path.exists(t_dir):
        continue
    for root, dirs, files in os.walk(t_dir):
        for file in files:
            path = os.path.join(root, file)
            if file.endswith(".xlsx"):
                res = search_excel(path)
                if res:
                    print(f"Found in Excel: {path} -> Sheet: {res[0]} -> Cell: {res[1]}")
                    found = True
            elif file.endswith(".docx"):
                res = search_docx(path)
                if res:
                    print(f"Found in Docx: {path} -> {res}")
                    found = True

if not found:
    print("NIJIMBERE not found in these target folders.")
