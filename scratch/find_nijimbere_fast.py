import os
import openpyxl
from docx import Document

workspace_dir = r"c:\Users\Degnon\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy"

def search_excel(path):
    try:
        # Ignore very large files
        if os.path.getsize(path) > 500000:
            return None
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell and "NIJIMBERE" in str(cell).upper():
                        return sheet, cell
    except Exception as e:
        pass
    return None

def search_docx(path):
    try:
        if os.path.getsize(path) > 500000:
            return None
        doc = Document(path)
        for p in doc.paragraphs:
            if "NIJIMBERE" in p.text.upper():
                return p.text
        for t in doc.tables:
            for r in t.rows:
                for cell in r.cells:
                    if "NIJIMBERE" in cell.text.upper():
                        return cell.text
    except Exception as e:
        pass
    return None

print("Scanning for NIJIMBERE...")
found = False
for root, dirs, files in os.walk(workspace_dir):
    if "node_modules" in root or ".git" in root or "MercyFiatMedSuiteDesktop" in root:
        continue
    for file in files:
        path = os.path.join(root, file)
        if "NIJIMBERE" in file.upper():
            print(f"MATCH IN FILENAME: {path}")
            found = True
        if file.endswith(".xlsx"):
            res = search_excel(path)
            if res:
                print(f"Found content in Excel: {path} -> Sheet: {res[0]} -> Cell: {res[1]}")
                found = True
        elif file.endswith(".docx"):
            res = search_docx(path)
            if res:
                print(f"Found content in Docx: {path} -> {res}")
                found = True

if not found:
    print("NIJIMBERE not found in any files.")
