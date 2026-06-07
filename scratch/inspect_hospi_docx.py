import docx
from docx.oxml.ns import qn

file_path = r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT HOSPI CMF\RAPPORT D'HOSPI CMF.docx"
doc = docx.Document(file_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Check textboxes
tb_count = 0
root = doc.element
for el in root.iter():
    if el.tag.endswith('txbxContent'):
        tb_count += 1

print(f"Total textboxes: {tb_count}")

# Print first 200 characters of first 10 paragraphs if any
print("\nFirst 5 paragraphs:")
for p in doc.paragraphs[:5]:
    if p.text.strip():
        print(f"  - {p.text[:100]}")

# Print some textbox content
print("\nFirst 10 textbox contents:")
tb_texts = []
count = 0
for el in root.iter():
    if el.tag.endswith('txbxContent'):
        for p in el.findall(qn('w:p')):
            p_obj = docx.text.paragraph.Paragraph(p, doc)
            text = p_obj.text.strip()
            if text:
                tb_texts.append(text)
                count += 1
                if count <= 15:
                    print(f"  TBox {count}: {text[:120]}")
        if count > 15:
            break
