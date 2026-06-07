import docx
import os

docx_path = r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT CONS\RAPPORT CS BATIA IBRAHIM.docx"
doc = docx.Document(docx_path)
print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")

# Print first 20 paragraphs even if they are empty
for idx in range(min(20, len(doc.paragraphs))):
    p = doc.paragraphs[idx]
    print(f"P {idx:02d}: len={len(p.text)} text={repr(p.text)}")

# Inspect first table if exists
if doc.tables:
    table = doc.tables[0]
    print(f"Table 0 rows: {len(table.rows)}")
    for r_idx, row in enumerate(table.rows[:5]):
        row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"Row {r_idx}: {row_text}")
