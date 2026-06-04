import docx
import os
import zipfile

f = r"c:\Users\Degnon\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\1. Document PC DR GIPSY\proforma\Dispense de sport AHAMADA.docx"
print("File size:", os.path.getsize(f))

# Check zip structure
try:
    with zipfile.ZipFile(f, 'r') as z:
        print("Zip contents:")
        for name in z.namelist()[:10]:
            print("  ", name)
except Exception as e:
    print("Error opening zip:", e)

try:
    doc = docx.Document(f)
    print("Paragraphs:", len(doc.paragraphs))
    print("Tables:", len(doc.tables))
    print("Sections:", len(doc.sections))
    # Check headers/footers
    for i, sec in enumerate(doc.sections):
        print(f"Section {i} header paragraphs:", len(sec.header.paragraphs))
        for p in sec.header.paragraphs:
            if p.text.strip():
                print("  Header p:", p.text)
        print(f"Section {i} footer paragraphs:", len(sec.footer.paragraphs))
        for p in sec.footer.paragraphs:
            if p.text.strip():
                print("  Footer p:", p.text)
                
    # Search elements
    from docx.oxml.ns import qn
    root = doc.element
    print("All tags inside element:")
    tags = set()
    for el in root.iter():
        tags.add(el.tag)
    print(sorted(list(tags))[:10])
except Exception as e:
    print("Docx error:", e)
