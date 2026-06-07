import docx
import os

files = [
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT HOSPI CMF\RAPPORT D'HOSPI CMF.docx",
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT CONS\RAPPORT DE CONSULTATION CMF.docx"
]

for f in files:
    if os.path.exists(f):
        doc = docx.Document(f)
        non_empty = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        print(f"File: {os.path.basename(f)}")
        print(f"  Total paragraphs: {len(doc.paragraphs)}")
        print(f"  Non-empty paragraphs: {len(non_empty)}")
        print(f"  First 10 non-empty paragraphs:")
        for idx, text in enumerate(non_empty[:10]):
            print(f"    {idx:02d}: {text[:100]}")
    else:
        print(f"Not found: {f}")
