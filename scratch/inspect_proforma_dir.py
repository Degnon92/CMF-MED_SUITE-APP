import os
import glob
import docx

workspace = r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy"
proforma_dir = os.path.join(workspace, "1. Document PC DR GIPSY", "proforma")

files = glob.glob(os.path.join(proforma_dir, "*.docx"))
print(f"Total docx files in proforma dir: {len(files)}")
print("Sample files:")
for f in files[:5]:
    print(" ", f)

for idx, f in enumerate(files[:3]):
    print(f"\n--- Reading content of file {idx}: {os.path.basename(f)} ---")
    try:
        doc = docx.Document(f)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        print(f"Paragraphs count: {len(paragraphs)}")
        print("First 10 paragraphs:")
        for p in paragraphs[:10]:
            print("  -", p)
        
        # Check for tables
        print(f"Tables count: {len(doc.tables)}")
        for t_idx, table in enumerate(doc.tables):
            print(f"Table {t_idx} (rows: {len(table.rows)}, cols: {len(table.columns)}):")
            for r_idx, row in enumerate(table.rows[:5]):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(f"  Row {r_idx}: {cells}")
    except Exception as e:
        print("Error reading file:", e)
