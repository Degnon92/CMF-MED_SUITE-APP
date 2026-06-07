import docx
import os
import glob

paths = [
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT CONS\*.docx",
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT HOSPI CMF\*.docx",
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\1. Document PC DR GIPSY\RAPPORT\*.docx"
]

found_text_docs = []
for p in paths:
    for filepath in glob.glob(p):
        try:
            doc = docx.Document(filepath)
            # count total non-empty chars
            total_chars = sum(len(para.text.strip()) for para in doc.paragraphs)
            if total_chars > 200:
                found_text_docs.append((filepath, total_chars, len(doc.tables)))
                if len(found_text_docs) >= 10:
                    break
        except Exception as e:
            pass
    if len(found_text_docs) >= 10:
        break

print(f"Found {len(found_text_docs)} text docs:")
for fd in found_text_docs:
    print(fd)
