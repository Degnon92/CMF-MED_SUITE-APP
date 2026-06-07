import os
import sys
import openpyxl
from docx import Document

workspace_dir = r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy"

def search_excel(path):
    try:
        # Ignore very large files
        if os.path.getsize(path) > 1000000:
            return None
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
                for c_idx, cell in enumerate(row):
                    if cell and "NIJIMBERE" in str(cell).upper():
                        return sheet, r_idx, c_idx, cell
    except Exception as e:
        pass
    return None

def search_docx(path):
    try:
        doc = Document(path)
        for p_idx, p in enumerate(doc.paragraphs):
            if "NIJIMBERE" in p.text.upper():
                return f"Paragraph {p_idx}: {p.text}"
        for t_idx, t in enumerate(doc.tables):
            for r_idx, r in enumerate(t.rows):
                for c_idx, cell in enumerate(r.cells):
                    if "NIJIMBERE" in cell.text.upper():
                        return f"Table {t_idx} Row {r_idx} Col {c_idx}: {cell.text}"
    except Exception as e:
        pass
    return None

print("Starting scan for NIJIMBERE...", flush=True)
found = False

for root, dirs, files in os.walk(workspace_dir):
    if "node_modules" in root or ".git" in root or "MercyFiatMedSuiteDesktop" in root:
        continue
    for file in files:
        path = os.path.join(root, file)
        
        # print progress for excel/docx files
        if file.endswith((".xlsx", ".docx")):
            print(f"Scanning {path}...", flush=True)
            
        if "NIJIMBERE" in file.upper():
            print(f"*** MATCH IN FILENAME: {path} ***", flush=True)
            found = True
            
        if file.endswith(".xlsx"):
            res = search_excel(path)
            if res:
                print(f"*** FOUND in Excel: {path} -> Sheet: {res[0]} -> Cell({res[1]},{res[2]}): {res[3]} ***", flush=True)
                found = True
        elif file.endswith(".docx"):
            res = search_docx(path)
            if res:
                print(f"*** FOUND in Docx: {path} -> {res} ***", flush=True)
                found = True

print("Scan completed.", flush=True)
