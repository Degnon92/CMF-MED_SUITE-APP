import docx
import os

files = [
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT HOSPI CMF\RAPPORT D'HOSPI CMF.docx",
    r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy\RAPPORT CONS\RAPPORT DE CONSULTATION CMF.docx"
]

for f in files:
    if os.path.exists(f):
        doc = docx.Document(f)
        print(f"File: {os.path.basename(f)}")
        print(f"  Total tables: {len(doc.tables)}")
        if doc.tables:
            for idx, table in enumerate(doc.tables[:5]):
                print(f"    Table {idx}: rows={len(table.rows)}")
                if table.rows:
                    first_row_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells[:4]]
                    print(f"      Row 0 cells: {first_row_cells}")
    else:
        print(f"Not found: {f}")
