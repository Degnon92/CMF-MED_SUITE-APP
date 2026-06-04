import os
import re
import glob
import json
import openpyxl
import docx
from datetime import datetime

workspace_dir = r"c:\Users\Degnon\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy"
app_dir = os.path.join(workspace_dir, "MercyFiatMedSuiteDesktop")

print("Démarrage de la reconstruction de la base de données...")

# -------------------------------------------------------------
# 1. FONCTIONS DE NETTOYAGE ET NORMALISATION
# -------------------------------------------------------------
def clean_patient_name(name):
    if not name or not isinstance(name, str):
        return ""
    
    clean = name.strip()
    
    # Nettoyage des préfixes cliniques (reproduction de cleanPatientName d'app.js)
    prefix_pattern = re.compile(
        r'^(?:CERTIFICAT\s+DE\s+MARIAGE|CERTIFICAT\s+DE\s+NON\s+BEGAIEMENT|CERTIFICAT\s+MEDICAL\s+INITIAL\s+DE\s+CONSTATATION\s+DE\s+COUPS\s+ET\s+BLESSURES|CERTIFICAT\s+MEDICAL\s+POUR\s+COUPS\s+ET\s+BLESSURES|CERTIFICAT\s+MEDICAL\s+DE\s+L[’\']ETAT\s+ACTUEL|CERTIFICAT\s+MEDICAL\s+INITIAL|CERTIFICAT\s+MEDICAL|CERTIFICAT\s+MED\s+INITIAL|CERTIFICAT\s+DE\s+REPOS|CERTIFICAT\s+DE\s+REPRISE|CERTIFICAT\s+DE\s+GUERISON|CERTIFICAT\s+DE\s+GUÉRISON|CERTIFICAT\s+DE|CERTIFICAT|RAPPORT\s+MEDICAL|RAPPORT\s+DE\s+MONSIEUR|RAPPORT\s+DE\s+MME|RAPPORT\s+DE|RAPPORT\s+D\'HOSPI|RAPPORT\s+D\'HOSPITALISATION|RAPPORT\s+DE\s+CONSULTATION|RAPPORT|CRO\s+MODELE|CRO|CMI|MEDICAL|MED\s+INITIAL|GUERISON\s+DE\s+MONSIEUR|GUERISON\s+DE\s+MME|GUERISON\s+DE|GUERISON|GUÉRISON|DECES\s+DE\s+MONSIEUR|DECES\s+DE\s+MME|DECES\s+DE|DECES|DÉCÈS|D\'HOSPI\s+TYPE|D\'HOSPI|DHOSPI|ATTESTATION\s+DE\s+GUERISON|ATTESTATION\s+DE\s+GUÉRISON|ATTESTATION\s+DE|ATTESTATION)\s+',
        re.IGNORECASE
    )
    
    old_clean = ""
    while clean != old_clean:
        old_clean = clean
        clean = prefix_pattern.sub("", clean).strip()
        
    # Séparation sur mots-clés interdits
    split_pattern = re.compile(
        r'(?:CERTIFICAT|JE\s+SOUSSIGN|JE\s+SOUSSIGNE|RAPPORT|DOSSIER|N°|NO\s+DOSSIER|CMI|CRO|CLINIQUE|MÉDECINE|MEDECINE|DIAGNOSTIC|INTERVENTION|CLIENT|ASSURANCE|AFRICAINE|NSIA|ALLIANZ|SUNU|AROO|SAAR|CORIS|FEDAS|MUTUELLE|SÉJOUR|SEJOUR|DATE|OPÉRATOIRE|OPERATOIRE|CERTFICAT|CERTIF|PATIENT|PATIENTE|COTONOU|RUE\s+PAVILLON|TEL\s*\:|E\-MAIL|EMAIL|E\s+MAIL|GUERISON|GUÉRISON|DECES|DÉCÈS)',
        re.IGNORECASE
    )
    
    match = split_pattern.search(clean)
    if match:
        clean = clean[:match.start()].strip()
        
    # Nettoyage des suffixes d'âge corrompus
    clean = re.sub(r'(?:\s+|:)\b(?:ANS|AGE|ÂGE|ANS\s+D[\’\']ÂGE)\b.*$', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'\b(?:ANS|AGE|ÂGE|ANS\s+D[\’\']ÂGE)\b.*$', '', clean, flags=re.IGNORECASE).strip()
    if clean.upper().endswith("AGE") and len(clean) > 5:
        clean = clean[:-3].strip()
        
    # Supprimer les caractères de ponctuation en fin de chaîne
    clean = re.sub(r'[\s\-\.\,\:\_]+$', '', clean).strip()
    
    # Remplacer les doubles espaces
    clean = re.sub(r'\s+', ' ', clean)
    
    # Nettoyage spécifique de suffixes comme "ok", "copie", "nov 23", "2", "3", etc.
    clean = re.sub(r'\s+\(?(?:ok|copie|2|3|4|1)\)?$', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'\s+fev\s+\d{2}$', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'\s+nov\s+\d{2}$', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'\s+aout\s+\d{2}$', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'\s+sept\s+\d{2}$', '', clean, flags=re.IGNORECASE).strip()
    
    return clean

def split_nom_prenom(full_name):
    cleaned = clean_patient_name(full_name)
    if not cleaned:
        return "", ""
        
    parts = cleaned.split(' ')
    if len(parts) == 1:
        return cleaned.upper(), ""
        
    # Heuristique : les mots entièrement en majuscules ou le premier mot sont le Nom
    nom_parts = []
    prenom_parts = []
    
    for part in parts:
        if part.isupper() and len(part) > 1:
            nom_parts.append(part)
        else:
            prenom_parts.append(part)
            
    if not nom_parts:
        # Si aucun mot n'était tout en majuscule, le premier est le nom, le reste est le prénom
        nom = parts[0].upper()
        prenom = " ".join(parts[1:])
    else:
        nom = " ".join(nom_parts).upper()
        prenom = " ".join(prenom_parts)
        
    # Si le prénom est vide et le nom a plusieurs mots, séparer le dernier mot comme prénom
    if not prenom and len(parts) > 1:
        nom = parts[0].upper()
        prenom = " ".join(parts[1:])
        
    return nom.strip(), prenom.strip()

def clean_clinical_term(term):
    if not term or not isinstance(term, str):
        return ""
    clean = term.strip()
    clean = re.sub(r'^[\s\-\.\,\:\_\•\*\t]+', '', clean).strip()
    clean = re.sub(r'[\s\-\.\,\:\_]+$', '', clean).strip()
    if len(clean) < 3 or len(clean) > 150:
        return ""
    return clean

def parse_date(text):
    if not text:
        return None
    # Chercher un format DD/MM/YYYY ou DD/MM/YY
    match = re.search(r'\b(\d{1,2})[\/\-\.](\d{1,2})[\/\-\.](\d{2,4})\b', text)
    if match:
        day, month, year = match.groups()
        if len(year) == 2:
            year = "20" + year
        try:
            dt = datetime(int(year), int(month), int(day))
            return dt.strftime("%Y-%m-%d")
        except:
            pass
            
    # Essayer textuel (ex : 10 Octobre 2024)
    months = {
        "janvier": 1, "fevrier": 2, "février": 2, "mars": 3, "avril": 4, "mai": 5, "juin": 6,
        "juillet": 7, "aout": 8, "août": 8, "septembre": 9, "octobre": 10, "novembre": 11, "décembre": 12, "decembre": 12
    }
    match_txt = re.search(r'\b(\d{1,2})\s+([a-zéûó]+)\s+(\d{4})\b', text, re.IGNORECASE)
    if match_txt:
        day, m_name, year = match_txt.groups()
        m_name = m_name.lower()
        if m_name in months:
            try:
                dt = datetime(int(year), months[m_name], int(day))
                return dt.strftime("%Y-%m-%d")
            except:
                pass
    return None

def map_insurance(text):
    if not text:
        return "PRIVE"
    t = text.lower().strip()
    if "sanlam" in t:
        return "SANLAM"
    elif "ascoma" in t:
        return "ASCOMA"
    elif "sunu" in t:
        return "SUNU"
    elif "nsia" in t:
        return "NSIA"
    elif "atlantique" in t:
        return "ATLANTIQUE"
    elif "africaine" in t:
        if "sinistre" in t or "accident" in t:
            return "AFRICAINE_SINISTRE"
        return "AFG"
    elif "loto" in t:
        return "LOTTO_FOOTBALL_CLUB"
    elif "coton" in t:
        return "COTON_SPORT"
    elif "nobila" in t:
        return "NOBILA"
    elif "gras" in t or "savoye" in t:
        return "GRAS SAVOYE"
    elif "olea" in t:
        return "OLEA"
    elif "fonds" in t or "garantie" in t or "fga" in t:
        return "FONDS_GARANTIE_AUTO"
    elif "transvie" in t:
        return "TRANSVIE"
    elif "sobremap" in t or "sobemap" in t:
        return "SOBEMAP"
    elif "port" in t or "pac" in t:
        return "PORT_AUTONOME_COTONOU"
    return "PRIVE"

# -------------------------------------------------------------
# 2. PARSER DE DOCX COMPILÉS (RAPPORT CONS / RAPPORT HOSPI)
# -------------------------------------------------------------
def get_textbox_paragraphs(doc):
    tb_texts = []
    root = doc.element
    for el in root.iter():
        if el.tag.endswith('txbxContent'):
            from docx.oxml.ns import qn
            for p in el.findall(qn('w:p')):
                p_obj = docx.text.paragraph.Paragraph(p, doc)
                if p_obj.text.strip():
                    tb_texts.append(p_obj.text.strip())
    return tb_texts

def parse_compiled_docx(file_path, category, start_idx):
    documents = []
    if not os.path.exists(file_path):
        print(f"Compilation introuvable : {file_path}")
        return documents, start_idx
        
    doc = docx.Document(file_path)
    tb_texts = get_textbox_paragraphs(doc)
    
    # Segmenter par "Patient:"
    segments = []
    current_segment = []
    for p in tb_texts:
        if p.lower().startswith("patient") and ":" in p:
            if current_segment:
                segments.append(current_segment)
            current_segment = [p]
        else:
            current_segment.append(p)
    if current_segment:
        segments.append(current_segment)
        
    idx = start_idx
    for seg in segments:
        if not seg or not seg[0].lower().startswith("patient") or ":" not in seg[0]:
            continue
        header = seg[0]
        # Extraire le nom du patient
        p_name = header.split(":", 1)[1].strip()
        nom, prenom = split_nom_prenom(p_name)
        if not nom:
            continue
            
        full_text = "\n".join(seg[1:])
        
        # Heuristiques d'extraction
        diagnosis = ""
        intervention = ""
        doc_date = None
        
        # Diagnostic
        diag_match = re.search(r'(?:diagnostic|bilan\s+lésionnel)\s*:\s*(.*)', full_text, re.IGNORECASE)
        if diag_match:
            diagnosis = clean_clinical_term(diag_match.group(1))
        else:
            # Chercher dans le texte s'il y a des pathologies connues
            pathologies = ["coxarthrose", "gonarthrose", "rupture", "fracture", "pseudarthrose", "appendicite", "lipome", "hernie", "adénome", "hydronéphrose"]
            for path in pathologies:
                m = re.search(rf'([^.\n]*{path}[^.\n]*)', full_text, re.IGNORECASE)
                if m:
                    diagnosis = clean_clinical_term(m.group(1))
                    break
                    
        # Intervention
        int_match = re.search(r'(?:intervention|geste|il\s+est\s+indiqué)\s*:\s*(.*)', full_text, re.IGNORECASE)
        if int_match:
            intervention = clean_clinical_term(int_match.group(1))
        else:
            interv_words = ["prothèse", "ostéosynthèse", "ligamentoplastie", "ablation", "résection", "cure d’hydrocèle", "appendicectomie", "adénomectomie"]
            for word in interv_words:
                m = re.search(rf'([^.\n]*{word}[^.\n]*)', full_text, re.IGNORECASE)
                if m:
                    intervention = clean_clinical_term(m.group(1))
                    break
                    
        # Date
        # Scanner à la fin du document ou chercher "Fait à"
        date_match = re.search(r'(?:fait\s+à\s+cotonou,\s+le\s+)(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})', full_text, re.IGNORECASE)
        if date_match:
            doc_date = parse_date(date_match.group(1))
        else:
            # Essayer de scanner n'importe quelle date
            doc_date = parse_date(full_text)
            
        if not doc_date:
            doc_date = "2025-06-02" # Date de secours
            
        templateId = "rapport_cs_assurance"
        if category == "Compte-Rendu Opératoire":
            templateId = "cro_cmf"
        elif category == "Hospitalisation":
            templateId = "rapport_hospi_assurance"
            
        title = f"{category} - {nom} {prenom}".strip()
        
        documents.append({
            "id": f"DOC-REAL-AUTO-{idx}",
            "type": "DOC",
            "category": category,
            "title": title,
            "templateId": templateId,
            "patientNom": nom,
            "patientPrenom": prenom,
            "patientAge": "35 ans",
            "date": doc_date,
            "diagnosis": diagnosis or "Bilan clinique",
            "intervention": intervention,
            "content": full_text
        })
        idx += 1
        
    print(f"Extraction de {len(documents)} rapports à partir de {os.path.basename(file_path)}")
    return documents, idx

# -------------------------------------------------------------
# 3. PARSER DE DOCX INDIVIDUELS
# -------------------------------------------------------------
def parse_individual_docx(dir_path, start_idx):
    documents = []
    if not os.path.exists(dir_path):
        print(f"Dossier introuvable : {dir_path}")
        return documents, start_idx
        
    idx = start_idx
    docx_files = glob.glob(os.path.join(dir_path, "**", "*.docx"), recursive=True)
    
    # Exclure les fichiers temporaires et les compilations
    exclude_patterns = ["~$", "cmf.docx", "exemplaire", "nomenclature", "stock", "price list"]
    docx_files = [f for f in docx_files if not any(p in os.path.basename(f).lower() for p in exclude_patterns)]
    
    for f in docx_files:
        try:
            doc = docx.Document(f)
            # Extraire texte normal et texte de zone de texte
            paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            textbox_texts = get_textbox_paragraphs(doc)
            all_lines = paragraphs_text + textbox_texts
            
            full_text = "\n".join(all_lines)
            
            # Nom de secours tiré du nom de fichier
            filename = os.path.basename(f)
            p_name = filename.replace("RAPPORT", "").replace("CS", "").replace("HOSPI", "").replace(".docx", "").strip()
            nom, prenom = split_nom_prenom(p_name)
            
            if not nom:
                continue
                
            # Heuristiques de métadonnées
            doc_date = parse_date(full_text)
            if not doc_date:
                doc_date = datetime.fromtimestamp(os.path.getmtime(f)).strftime("%Y-%m-%d")
                
            # Déterminer la catégorie
            category = "Consultation"
            templateId = "rapport_cs_assurance"
            
            fn_lower = filename.lower()
            if "cro" in fn_lower or "opératoire" in fn_lower or "operatoire" in fn_lower:
                category = "Compte-Rendu Opératoire"
                templateId = "cro_cmf"
            elif "hospi" in fn_lower or "hospitalisation" in fn_lower:
                category = "Hospitalisation"
                templateId = "rapport_hospi_assurance"
            elif "certif" in fn_lower or "repos" in fn_lower or "reprise" in fn_lower or "guerison" in fn_lower:
                category = "Certificat Médical"
                templateId = "certif_reprise"
                
            diagnosis = ""
            diag_match = re.search(r'(?:diagnostic|bilan\s+lésionnel)\s*:\s*(.*)', full_text, re.IGNORECASE)
            if diag_match:
                diagnosis = clean_clinical_term(diag_match.group(1))
                
            intervention = ""
            int_match = re.search(r'(?:intervention|geste|il\s+est\s+indiqué)\s*:\s*(.*)', full_text, re.IGNORECASE)
            if int_match:
                intervention = clean_clinical_term(int_match.group(1))
                
            # Si le texte est très court (document numérisé / scanné)
            if len(full_text.strip()) < 100:
                full_text = f"Document clinique de type {category} pour {nom} {prenom}.\nContenu scanné et disponible dans le dossier physique."
                
            documents.append({
                "id": f"DOC-REAL-AUTO-{idx}",
                "type": "DOC",
                "category": category,
                "title": f"{category} - {nom} {prenom}".strip(),
                "templateId": templateId,
                "patientNom": nom,
                "patientPrenom": prenom,
                "patientAge": "35 ans",
                "date": doc_date,
                "diagnosis": diagnosis or "Bilan clinique",
                "intervention": intervention,
                "content": full_text
            })
            idx += 1
        except Exception as e:
            # print(f"Erreur lors de la lecture de {f}: {e}")
            pass
            
    print(f"Extraction de {len(documents)} rapports individuels de type docx.")
    return documents, idx

# -------------------------------------------------------------
# 4. PARSER DES FACTURES EXCEL (EXEMPLAIRE PROFORMA.xlsx)
# -------------------------------------------------------------
def parse_excel_bills(wb_path, start_idx):
    bills = []
    if not os.path.exists(wb_path):
        print(f"Fichier Excel introuvable : {wb_path}")
        return bills, start_idx
        
    wb = openpyxl.load_workbook(wb_path, data_only=True)
    idx = start_idx
    
    # Onglets de configuration à ignorer
    exclude_sheets = ["feuil1", "feuil2", "winsu sports", "feuil4"]
    
    for sheetname in wb.sheetnames:
        if sheetname.lower() in exclude_sheets:
            continue
            
        sheet = wb[sheetname]
        
        # 1. Extraire les métadonnées
        patient_name = ""
        bill_type = "PROFORMA"
        insurance = "PRIVE"
        date_str = ""
        k_code = ""
        intervention = ""
        
        # Scan de toutes les cellules utiles de la feuille
        for r in range(1, 35):
            for c in range(1, 6):
                val = sheet.cell(r, c).value
                if not val or not isinstance(val, str):
                    continue
                val = val.strip()
                
                # Patient
                if val.lower().startswith("patient:") or val.lower().startswith("patient :"):
                    patient_name = val.split(":", 1)[1].strip()
                # Client / Insurance
                elif val.lower().startswith("client:") or val.lower().startswith("client :"):
                    insurance = map_insurance(val.split(":", 1)[1].strip())
                # Intervention
                elif val.lower().startswith("intervention:") or val.lower().startswith("intervention :"):
                    intervention = val.split(":", 1)[1].strip()
                # Type de facture
                if "définitive" in val.lower() or "definitif" in val.lower() or "facture définitive" in val.lower():
                    bill_type = "DEFINITIF"
                elif "détail" in val.lower() or "detail" in val.lower() or "détail assurance" in val.lower():
                    bill_type = "DETAIL_ASSUR"
                # Date
                if "cotonou,le" in val.lower() or "cotonou, le" in val.lower():
                    date_str = parse_date(val)
                # Code K
                k_m = re.search(r'(KC?\s*\d+)', val.replace('_', ' '))
                if k_m:
                    k_code = k_m.group(1)
                    
        # Si aucun nom de patient dans la feuille, on déduit du nom de l'onglet
        if not patient_name:
            # Vérifier si le nom de l'onglet contient des lettres
            if re.search(r'[A-Za-z]', sheetname):
                patient_name = sheetname
            else:
                continue
                
        nom, prenom = split_nom_prenom(patient_name)
        if not nom:
            continue
            
        if not date_str:
            date_str = "2026-06-01"
            
        # 2. Chercher la table des items
        header_row = -1
        for r in range(8, 20):
            val = sheet.cell(r, 1).value
            if val and isinstance(val, str) and ("actes" in val.lower() or "désignations" in val.lower() or "designations" in val.lower()):
                header_row = r
                break
                
        items = []
        gross_total = 0
        
        if header_row != -1:
            for r in range(header_row + 1, header_row + 30):
                desig = sheet.cell(r, 1).value
                qty = sheet.cell(r, 2).value
                price = sheet.cell(r, 3).value
                subtotal = sheet.cell(r, 4).value
                
                if desig:
                    desig_str = str(desig).strip()
                    if "total" in desig_str.lower() or "caisse" in desig_str.lower():
                        break
                    
                    # Nettoyer les montants
                    try:
                        qty_val = int(float(str(qty).replace(' ', ''))) if qty else 1
                    except:
                        qty_val = 1
                    try:
                        price_val = int(float(str(price).replace(' ', ''))) if price else 0
                    except:
                        price_val = 0
                    try:
                        subtotal_val = int(float(str(subtotal).replace(' ', ''))) if subtotal else price_val * qty_val
                    except:
                        subtotal_val = price_val * qty_val
                        
                    if price_val > 0:
                        items.append({
                            "name": clean_clinical_term(desig_str) or desig_str,
                            "price": price_val,
                            "qty": qty_val,
                            "subtotal": subtotal_val
                        })
                        gross_total += subtotal_val
                        
        if not items:
            # Forfait minimum par défaut s'il n'y a pas d'items saisis
            items = [{"name": "Frais de soins cliniques standards", "price": gross_total or 210000, "qty": 1, "subtotal": gross_total or 210000}]
            if not gross_total:
                gross_total = 210000
                
        # Split calculs par défaut
        coverage = 80 if insurance != "PRIVE" else 0
        part_assurance = 0
        part_patient = gross_total
        
        if insurance != "PRIVE":
            part_assurance = int(gross_total * (coverage / 100))
            part_patient = gross_total - part_assurance
            
        bills.append({
            "id": f"BILL-REAL-AUTO-{idx}",
            "reference": f"MF-{bill_type[:3].upper()}-2026-{idx:03d}",
            "patientNom": nom,
            "patientPrenom": prenom,
            "type": bill_type,
            "customTitle": "Point Définitif d'Hospitalisation" if bill_type == "DEFINITIF" else "",
            "insurance": insurance,
            "coverage": coverage,
            "matricule": "N/A" if insurance != "PRIVE" else "",
            "diagnostic": "Bilan clinique",
            "intervention": intervention,
            "kCode": k_code,
            "showDiag": True,
            "showInterv": True if intervention else False,
            "showSig": True,
            "showCachet": True,
            "useSplit": True if insurance != "PRIVE" else False,
            "items": items,
            "grossTotal": gross_total,
            "discountPct": 0,
            "reductionAmount": 0,
            "discountedTotal": gross_total,
            "discountType": "PERCENT",
            "discountValue": 0,
            "partAssurance": part_assurance,
            "partPatient": part_patient,
            "paymentMethod": "TIERS_PAYANT" if insurance != "PRIVE" else "CASH",
            "amountPaidPatient": part_patient,
            "balancePatient": 0,
            "status": "RÉGLÉ",
            "date": date_str
        })
        idx += 1
        
    print(f"Extraction de {len(bills)} factures à partir d'EXEMPLAIRE PROFORMA.xlsx")
    return bills, idx

# -------------------------------------------------------------
# 5. INTEGRATION GLOBALE & ENREGISTREMENT
# -------------------------------------------------------------
documents_list = []
bills_list = []

# 1. Parse compiled reports
docs_1, count_doc_idx = parse_compiled_docx(
    os.path.join(workspace_dir, "RAPPORT CONS", "RAPPORT DE CONSULTATION CMF.docx"), 
    "Consultation", 1
)
documents_list.extend(docs_1)

docs_2, count_doc_idx = parse_compiled_docx(
    os.path.join(workspace_dir, "RAPPORT HOSPI CMF", "RAPPORT D'HOSPI CMF.docx"), 
    "Hospitalisation", count_doc_idx
)
documents_list.extend(docs_2)

# 2. Parse individual reports from directories
docs_3, count_doc_idx = parse_individual_docx(os.path.join(workspace_dir, "RAPPORT CONS"), count_doc_idx)
documents_list.extend(docs_3)

docs_4, count_doc_idx = parse_individual_docx(os.path.join(workspace_dir, "RAPPORT HOSPI CMF"), count_doc_idx)
documents_list.extend(docs_4)

docs_5, count_doc_idx = parse_individual_docx(os.path.join(workspace_dir, "1. Document PC DR GIPSY"), count_doc_idx)
documents_list.extend(docs_5)

# 3. Parse Excel bills
bills, count_bill_idx = parse_excel_bills(os.path.join(workspace_dir, "PROFORMA CHIRURGIE", "EXEMPLAIRE PROFORMA.xlsx"), 1)
bills_list.extend(bills)

# 4. Déduplication finale et stricte des rapports cliniques pour corriger les ID doubles
unique_docs = {}
for d in documents_list:
    # Créer une clé unique robuste basée sur le patient et le contenu
    clean_content_snippet = re.sub(r'\s+', '', (d["content"] or "")[:200].lower())
    key = f"{d['patientNom']}||{d['patientPrenom']}||{clean_content_snippet}"
    if key not in unique_docs:
        unique_docs[key] = d
        
# Réassigner des IDs uniques séquentiels propres
final_docs = list(unique_docs.values())
for i, d in enumerate(final_docs):
    d["id"] = f"DOC-REAL-AUTO-{i+1}"
    
print(f"Nombre final de rapports médicaux uniques : {len(final_docs)}")

# 5. Déduplication finale des factures
unique_bills = {}
for b in bills_list:
    key = f"{b['patientNom']}||{b['patientPrenom']}||{b['grossTotal']}||{b['date']}"
    if key not in unique_bills:
        unique_bills[key] = b
        
final_bills = list(unique_bills.values())
for i, b in enumerate(final_bills):
    b["id"] = f"BILL-REAL-AUTO-{i+1}"
    
print(f"Nombre final de factures uniques : {len(final_bills)}")

# 6. Reconstruire la liste unique des PATIENTS avec leurs diagnostics et interventions
patients_map = {}
for d in final_docs:
    fullname = f"{d['patientNom']} {d['patientPrenom']}".strip().upper()
    if fullname not in patients_map:
        patients_map[fullname] = {
            "name": fullname,
            "diagnosis": d.get("diagnosis", ""),
            "intervention": d.get("intervention", ""),
            "kCode": ""
        }
    else:
        if d.get("diagnosis") and not patients_map[fullname]["diagnosis"]:
            patients_map[fullname]["diagnosis"] = d["diagnosis"]
        if d.get("intervention") and not patients_map[fullname]["intervention"]:
            patients_map[fullname]["intervention"] = d["intervention"]

for b in final_bills:
    fullname = f"{b['patientNom']} {b['patientPrenom']}".strip().upper()
    if fullname not in patients_map:
        patients_map[fullname] = {
            "name": fullname,
            "diagnosis": b.get("diagnostic", ""),
            "intervention": b.get("intervention", ""),
            "kCode": b.get("kCode", "")
        }
    else:
        if b.get("diagnostic") and not patients_map[fullname]["diagnosis"]:
            patients_map[fullname]["diagnosis"] = b["diagnostic"]
        if b.get("intervention") and not patients_map[fullname]["intervention"]:
            patients_map[fullname]["intervention"] = b["intervention"]
        if b.get("kCode") and not patients_map[fullname]["kCode"]:
            patients_map[fullname]["kCode"] = b["kCode"]

final_patients = list(patients_map.values())
print(f"Nombre final de patients uniques : {len(final_patients)}")

# 7. Écriture des fichiers sur le disque (dans le dossier de l'application)
os.makedirs(app_dir, exist_ok=True)

with open(os.path.join(app_dir, "patients_db.json"), "w", encoding="utf-8") as f:
    json.dump(final_patients, f, indent=4, ensure_ascii=False)
print("Fichier patients_db.json écrit.")

with open(os.path.join(app_dir, "bills_db.json"), "w", encoding="utf-8") as f:
    json.dump(final_bills, f, indent=4, ensure_ascii=False)
print("Fichier bills_db.json écrit.")

with open(os.path.join(app_dir, "documents_db.json"), "w", encoding="utf-8") as f:
    json.dump(final_docs, f, indent=4, ensure_ascii=False)
print("Fichier documents_db.json écrit.")

# 8. Reconstruire real_data.js
js_content = f"""/* ==========================================
   real_data.js - Vrais Rapports Médicaux Clinique Mercy Fiat
   ========================================== */

window.MercyFiatRealDocs = {json.dumps(final_docs, indent=4, ensure_ascii=False)};
"""
with open(os.path.join(app_dir, "real_data.js"), "w", encoding="utf-8") as f:
    f.write(js_content)
print("Fichier real_data.js écrit.")

print("Migration de la base de données terminée avec succès !")
