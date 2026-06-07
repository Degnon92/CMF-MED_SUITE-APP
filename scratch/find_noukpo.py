import os
import openpyxl
from docx import Document

workspace_dir = r"c:\Users\Farus\Documents\2.MERCY FIAT CLINIQUE\2. Dr Gipsy"

def search_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)
        if "NOUKPOZOUNKOU" in text.upper():
            return text
    except Exception as e:
        pass
    return None

def search_xlsx(file_path):
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        for sheetname in wb.sheetnames:
            if "NOUKPOZOUNKOU" in sheetname.upper():
                return f"Sheet name: {sheetname}"
            # Also search cells if needed
    except Exception as e:
        pass
    return None

for root, dirs, files in os.walk(workspace_dir):
    for file in files:
        file_path = os.path.join(root, file)
        if file.endswith('.docx'):
            res = search_docx(file_path)
            if res:
                print(f"Found in DOCX: {file_path}")
                print(res[:500] + "...")
        elif file.endswith('.xlsx'):
            res = search_xlsx(file_path)
            if res:
                print(f"Found in XLSX: {file_path} - {res}")
