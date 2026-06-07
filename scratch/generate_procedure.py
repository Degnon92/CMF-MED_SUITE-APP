import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

def main():
    doc_path = r"c:\Users\Farus\Pictures\Screenshots\AA\procedure_prise_en_charge.docx"
    pdf_path = r"c:\Users\Farus\Pictures\Screenshots\AA\procedure_prise_en_charge.pdf"
    img_dir = r"c:\Users\Farus\Pictures\Screenshots\AA"

    doc = Document()

    # Define color scheme (Forest Green #1B4D3E and Dark Charcoal #333333)
    primary_color = RGBColor(27, 77, 62)
    secondary_color = RGBColor(47, 79, 79)
    text_color = RGBColor(51, 51, 51)

    # Set document margins (0.75 inches for a modern look)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Configure the 'Normal' body style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = text_color

    # Add a main title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("PROCÉDURE DE PRISE EN CHARGE DES PATIENTS")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = primary_color

    # Add a subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(24)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Portail TPA (Orsys) — L'AFRICAINE DES ASSURANCES BENIN\nClinique Mercy Fiat")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = secondary_color

    # Add introduction
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(18)
    intro_run = intro_p.add_run(
        "Ce guide pratique décrit, étape par étape, la procédure à suivre sur le portail TPA pour l'enregistrement "
        "des consultations et la saisie des ordonnances de pharmacie pour les patients assurés auprès de l'Afrique des Assurances."
    )
    intro_run.font.name = 'Arial'
    intro_run.font.size = Pt(10.5)

    # Steps list: (image_filename, title, instruction_list)
    steps = [
        ("1. CONSULTATION.png", "Étape 1 : Accès à l'identification de l'assuré", [
            "Sur le tableau de bord principal du portail Tiers payant (TPA), cliquez sur le module \"Identification Assuré\" (encadré en orange sur l'écran)."
        ]),
        ("2. nUMERO.png", "Étape 2 : Recherche de l'assuré", [
            "Saisissez le code de la carte de l'assuré dans le champ de recherche \"Inscrivez le code de la carte\".",
            "Cliquez sur le bouton \"Identifier\" pour valider et afficher les informations de l'assuré."
        ]),
        ("3. CLIQUE PRESTATION.png", "Étape 3 : Initialisation de la prestation", [
            "Vérifiez le statut de l'adhérent (celui-ci doit être marqué \"Couvert\" en vert).",
            "Cliquez sur le bouton jaune \"Faire Une Prestation\" afin d'initier la feuille de soins."
        ]),
        ("4. CHOIX AFFECTION.png", "Étape 4 : Sélection de la famille de maladie (Affection)", [
            "Dans la fenêtre d'ajout de feuille de soins, sélectionnez l'affection correspondante dans le champ \"Famille Maladie\".",
            "Cliquez ensuite sur le bouton bleu \"+ Ajouter un Soin\" pour passer à l'enregistrement de l'acte."
        ]),
        ("5.sAUVEGARDER cONSULTATION.png", "Étape 5 : Enregistrement de l'acte de consultation", [
            "Remplissez les détails du soin : sélectionnez la \"Prestation\", rédigiez les observations nécessaires dans \"Observation prestataire\", puis activez \"Impression Observation?\" si souhaité.",
            "Renseignez le \"Coût unitaire\" (débours réel), puis cliquez sur le bouton bleu \"Sauvegarder\"."
        ]),
        ("6. lcLIQUEZ Sur fermer.png", "Étape 6 : Validation et fermeture de la consultation", [
            "Après avoir enregistré l'acte, celui-ci s'affiche dans le récapitulatif des soins de la feuille.",
            "Cliquez sur le bouton rouge \"Fermer\" en haut à droite pour revenir à la liste générale."
        ]),
        ("7. accueil ou ouverture barre latérale pour gestion feuille soin.png", "Étape 7 : Déploiement de la barre latérale", [
            "Cliquez sur la flèche d'extension de la barre latérale ( \">\" ) située dans le coin inférieur gauche pour afficher le menu principal."
        ]),
        ("8. recherche patient et cliquer sur flècje.png", "Étape 8 : Recherche et modification de la feuille de soins", [
            "Dans le menu latéral, cliquez sur \"Gestion Feuilles de soins\".",
            "Retrouvez la ligne du patient concerné (statut de la feuille : \"Ouverte\").",
            "Cliquez sur l'icône crayon orange (Modifier) à droite de la ligne pour ouvrir à nouveau la feuille de soins."
        ]),
        ("8. Pharmacie.png", "Étape 9 : Sélection de l'onglet Pharmacie", [
            "Dans la feuille de soins en mode édition, cliquez sur l'onglet \"Pharmacie\"."
        ]),
        ("9. Ajouter med.png", "Étape 10 : Ajout d'une ordonnance", [
            "Sous l'onglet Pharmacie, cliquez sur le bouton bleu \"+ Ajouter ordonnance\"."
        ]),
        ("10. A reprendre.png", "Étape 11 : Saisie de l'ordonnance (Attention : Bug d'affichage fréquent)", [
            "Attention / Dysfonctionnement fréquent : La première fois que vous ouvrez cette fenêtre, la liste des médicaments ne s'affiche pas et le champ reste inactif (représenté par la croix rouge sur l'image).",
            "Procédure de contournement :",
            "  1. Cliquez sur le bouton rouge \"Fermer\" en haut à droite du pop-up pour fermer la fenêtre de l'ordonnance.",
            "  2. Cliquez de nouveau sur le bouton bleu \"+ Ajouter ordonnance\" (répéter l'Étape 10)."
        ]),
        ("11. Bon taper medicament.png", "Étape 12 : Saisie de l'ordonnance (Recherche de médicaments active)", [
            "Une fois la fenêtre rouverte, cliquez à nouveau dans le premier champ \"Désignation\".",
            "Vérifiez qu'une petite icône de chargement rotative (composée de 4 points rouges) s'affiche et tourne dans le champ de saisie (indiquée par la flèche rouge).",
            "Dès que cette icône tourne, saisissez le nom du médicament pour voir s'afficher la liste de suggestions de la base de données, puis sélectionnez-le.",
            "Renseignez le champ \"Nombre\" (quantité) et les observations médicales si nécessaire, puis cliquez sur le bouton bleu \"Sauvegarder\"."
        ])
    ]

    for idx, (filename, step_title, instructions) in enumerate(steps, 1):
        # Add heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        run_title = h.add_run(step_title)
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(13)
        run_title.font.bold = True
        run_title.font.color.rgb = primary_color

        # Add image
        img_path = os.path.join(img_dir, filename)
        if os.path.exists(img_path):
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.space_after = Pt(8)
            # Standard page width is 8.5", margins 0.75" * 2 = 1.5", printable width = 7.0".
            # We scale the images to 6.2 inches to fit nicely.
            img_p.add_run().add_picture(img_path, width=Inches(6.2))
        else:
            doc.add_paragraph(f"[Image manquante : {filename}]")

        # Add instructions
        for instr in instructions:
            inst_p = doc.add_paragraph(style='List Bullet')
            inst_p.paragraph_format.space_after = Pt(3)
            inst_p.paragraph_format.left_indent = Inches(0.25)
            
            run_instr = inst_p.add_run(instr)
            run_instr.font.name = 'Arial'
            run_instr.font.size = Pt(10)
            
            # Format "Attention" or "Important" in bold
            if "Attention" in instr or "Important" in instr:
                run_instr.font.bold = True
                run_instr.font.color.rgb = RGBColor(180, 0, 0)

        # Add space after step
        space_p = doc.add_paragraph()
        space_p.paragraph_format.space_after = Pt(12)

    # Save Word document
    print(f"Sauvegarde du document Word dans : {doc_path}")
    doc.save(doc_path)

    # Convert to PDF
    print(f"Conversion en PDF : {pdf_path}")
    convert(doc_path, pdf_path)
    print("Conversion terminée avec succès !")

if __name__ == '__main__':
    main()
